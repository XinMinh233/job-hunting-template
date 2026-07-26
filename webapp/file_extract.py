from __future__ import annotations

import argparse
import resource
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}
MAX_EXTRACTED_CHARS = 2_000_000
MAX_DOCX_EXPANDED = 50 * 1024 * 1024


def extract_text(path: Path, suffix: str) -> str | None:
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")[
            :MAX_EXTRACTED_CHARS
        ]
    if suffix == ".pdf":
        reader = PdfReader(str(path), strict=False)
        if len(reader.pages) > 200:
            raise ValueError("PDF 页数超过 200 页")
        parts: list[str] = []
        total = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            parts.append(text)
            total += len(text)
            if total > MAX_EXTRACTED_CHARS:
                break
        return "\n\n".join(parts)[:MAX_EXTRACTED_CHARS]
    if suffix == ".docx":
        with zipfile.ZipFile(path) as archive:
            if len(archive.infolist()) > 10_000:
                raise ValueError("DOCX 内部文件数量异常")
            expanded = sum(item.file_size for item in archive.infolist())
            if expanded > MAX_DOCX_EXPANDED:
                raise ValueError("DOCX 解压后体积过大")
            if any(
                item.file_size > 0
                and (
                    item.compress_size == 0
                    or item.file_size / item.compress_size > 200
                )
                for item in archive.infolist()
            ):
                raise ValueError("DOCX 压缩比例异常")
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)[:MAX_EXTRACTED_CHARS]
    return None


def _apply_limits() -> None:
    limits = (
        (resource.RLIMIT_CPU, 20),
        (resource.RLIMIT_AS, 512 * 1024 * 1024),
        (resource.RLIMIT_FSIZE, 12 * 1024 * 1024),
        (resource.RLIMIT_NOFILE, 64),
    )
    for limit, value in limits:
        try:
            resource.setrlimit(limit, (value, value))
        except (OSError, ValueError):
            pass


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("suffix", choices=(".pdf", ".docx"))
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    _apply_limits()
    text = extract_text(args.source, args.suffix)
    args.output.write_text(text or "", encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

